"""
건축 공모 & 법규 분석 시스템 v4.0 - Professional Edition
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 주요 기능:
- 공모지침서 분석 (용도, 지역지구, 설계조건)
- 법규 위계 분석 (상위법 vs 하위법)
- 교차 분석 (지침 + 법규 통합)
- 실별 면적표 시각화 (Plotly)
- 전문 보고서 생성 (한글 양식)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

필수 라이브러리 설치:
pip install streamlit google-generativeai python-dotenv python-docx plotly pandas
"""

import streamlit as st
import google.generativeai as genai
import os
import time
import tempfile
from pathlib import Path
from dotenv import load_dotenv
from datetime import datetime
from io import BytesIO
import json
import re

# 데이터 처리
import pandas as pd

# 그래프
import plotly.express as px
import plotly.graph_objects as go

# 문서 생성
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# .env 파일 로드
load_dotenv()

# ================================
# 페이지 설정
# ================================
st.set_page_config(
    page_title="건축 공모 & 법규 분석 시스템 v4.0",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================================
# 저작권 문구
# ================================
st.markdown("""
<div style='text-align: right; color: #999; font-size: 0.75rem; padding: 0.5rem;'>
    All intellectual property rights belong to Kim Doyoung.
</div>
""", unsafe_allow_html=True)

# ================================
# 고급 커스텀 CSS
# ================================
st.markdown("""
<style>
    /* 메인 타이틀 */
    .main-title {
        text-align: center;
        background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%);
        color: white;
        padding: 2rem;
        border-radius: 15px;
        font-size: 2.5rem;
        font-weight: bold;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    /* 버전 뱃지 */
    .version-badge {
        display: inline-block;
        background: #f59e0b;
        color: white;
        padding: 0.3rem 1rem;
        border-radius: 20px;
        font-size: 0.9rem;
        margin-left: 1rem;
    }
    
    /* 분석 버튼 */
    .stButton > button {
        width: 100%;
        background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%);
        color: white;
        font-size: 1.2rem;
        font-weight: bold;
        padding: 1rem;
        border-radius: 12px;
        border: none;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(0,0,0,0.15);
    }
    
    /* 섹션 헤더 */
    .section-header {
        background: linear-gradient(135deg, #f3f4f6 0%, #e5e7eb 100%);
        padding: 1rem;
        border-left: 5px solid #3b82f6;
        border-radius: 8px;
        margin: 1rem 0;
    }
    
    /* 결과 카드 */
    .result-card {
        background: white;
        border: 2px solid #e5e7eb;
        border-radius: 10px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
    }
</style>
""", unsafe_allow_html=True)

# ================================
# 사이드바
# ================================
with st.sidebar:
    st.markdown("## ⚙️ 설정")
    
    st.markdown("### 🔑 API 인증")
    
    env_api_key = os.getenv("GOOGLE_API_KEY", "")
    
    if env_api_key:
        st.success("✅ API 키 로드 완료!")
        api_key = env_api_key
    else:
        st.warning("⚠️ API 키를 입력하세요")
        api_key = st.text_input(
            "Google Gemini API Key",
            type="password",
            help="https://aistudio.google.com/app/apikey"
        )
    
    if api_key:
        genai.configure(api_key=api_key)
        st.success("🎯 API 연결 완료!")
    
    st.markdown("---")
    
    # 모델 선택
    st.markdown("### 🤖 AI 모델")
    selected_model = "models/gemini-2.5-flash"
    st.info(f"✅ {selected_model}")
    
    st.markdown("---")
    
    # 분석 옵션
    st.markdown("### 🎛️ 분석 옵션")
    
    analysis_depth = st.selectbox(
        "분석 상세도",
        ["표준", "상세", "매우 상세"],
        index=1
    )
    
    include_visualization = st.checkbox(
        "📊 실별 면적표 시각화",
        value=True
    )
    
    st.markdown("---")
    st.markdown("""
    ### 📚 v4.0 신기능
    - 🏛️ 공모지침서 분석
    - ⚖️ 법규 위계 분석
    - 🔄 교차 분석
    - 📊 면적표 시각화
    """)


# ================================
# 핵심 함수: PDF 업로드
# ================================

def upload_pdf_to_gemini(uploaded_file, display_name=None):
    """단일 PDF 파일을 Gemini에 업로드"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        name = display_name or uploaded_file.name
        uploaded_gemini_file = genai.upload_file(tmp_path, display_name=name)
        
        # 처리 대기
        while uploaded_gemini_file.state.name == "PROCESSING":
            time.sleep(1)
            uploaded_gemini_file = genai.get_file(uploaded_gemini_file.name)
        
        os.unlink(tmp_path)
        
        return uploaded_gemini_file if uploaded_gemini_file.state.name != "FAILED" else None
        
    except Exception as e:
        st.error(f"❌ 업로드 오류: {str(e)}")
        return None


def upload_multiple_pdfs(uploaded_files, prefix="법규"):
    """여러 PDF 파일을 Gemini에 업로드"""
    uploaded_files_list = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for idx, file in enumerate(uploaded_files, 1):
        status_text.markdown(f"📤 {prefix} {idx}/{len(uploaded_files)}: {file.name} 업로드 중...")
        
        gemini_file = upload_pdf_to_gemini(file, f"{prefix}_{idx}")
        
        if gemini_file:
            uploaded_files_list.append(gemini_file)
            status_text.markdown(f"✅ {file.name} 완료!")
        
        progress_bar.progress(idx / len(uploaded_files))
        time.sleep(0.3)
    
    status_text.markdown(f"🎉 전체 {len(uploaded_files_list)}/{len(uploaded_files)} 파일 업로드 완료!")
    
    return uploaded_files_list


# ================================
# 핵심 함수: 공모지침서 분석
# ================================

def analyze_competition_guidelines(gemini_file, model_name):
    """
    공모지침서 분석
    - 용도, 지역지구, 설계조건 파악
    - 실별 면적표 추출
    """
    
    prompt = """
당신은 건축 공모 전문가입니다.
첨부된 공모지침서를 분석하여 다음 정보를 JSON 형식으로 추출하세요.

**추출할 정보:**
1. 프로젝트 개요
   - 사업명
   - 위치 (주소)
   - 지역지구
   - 건축물 용도

2. 설계 조건
   - 대지면적
   - 건축면적
   - 건폐율 (%)
   - 용적률 (%)
   - 층수 제한
   - 높이 제한

3. 실별 면적표 (있는 경우)
   - 실명과 면적을 배열로
   - 예: [{"실명": "로비", "면적": 100}, ...]

4. 전용/공용면적 (있는 경우)
   - 전용면적 합계
   - 공용면적 합계

**출력 형식 (JSON):**
```json
{
  "프로젝트명": "...",
  "위치": "...",
  "지역지구": "...",
  "용도": "...",
  "대지면적": "...",
  "건폐율": "...",
  "용적률": "...",
  "층수제한": "...",
  "실별면적표": [...],
  "전용면적": "...",
  "공용면적": "..."
}
```

**중요:**
- JSON 형식만 출력하세요 (추가 설명 없이)
- 정보가 없으면 빈 문자열 또는 빈 배열로
- 숫자는 문자열로 (예: "60%")
"""
    
    try:
        st.info("🔍 공모지침서 분석 중...")
        
        model = genai.GenerativeModel(model_name)
        response = model.generate_content([gemini_file, prompt])
        
        # JSON 파싱
        result_text = response.text
        
        # JSON 추출 (```json 제거)
        json_match = re.search(r'```json\s*(.*?)\s*```', result_text, re.DOTALL)
        if json_match:
            result_text = json_match.group(1)
        
        result_data = json.loads(result_text)
        
        st.success("✅ 공모지침서 분석 완료!")
        
        return result_data
        
    except Exception as e:
        st.error(f"❌ 공모지침서 분석 오류: {str(e)}")
        return None


# ================================
# 핵심 함수: 법규 분석 (위계 포함)
# ================================

def analyze_regulations_with_hierarchy(gemini_files, guideline_data, model_name):
    """
    법규 분석 (상위법/하위법 구분)
    - 국계법과 조례 구분
    - 실질 적용 법규 강조
    """
    
    # 공모 정보 추출
    location = guideline_data.get("위치", "")
    zone = guideline_data.get("지역지구", "")
    usage = guideline_data.get("용도", "")
    
    prompt = f"""
당신은 대한민국 건축법 전문가입니다.
첨부된 법규 PDF 문서들을 분석하여, 아래 프로젝트에 적용되는 법규를 **위계별로** 정리하세요.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📍 **프로젝트 정보**
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
• 위치: {location}
• 지역지구: {zone}
• 용도: {usage}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚖️ **법규 위계 분석 원칙**
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. **상위법 (국계법)**
   - 건축법, 국토의 계획 및 이용에 관한 법률 등
   - 전국 공통 적용
   
2. **하위법 (조례)**
   - OO시 건축조례, OO시 도시계획조례 등
   - 지역 특화 기준
   
3. **실질 적용 원칙**
   - 상위법이 조례로 위임한 경우 → 조례 기준 우선
   - 조례가 상위법보다 엄격한 경우 → 조례 기준 적용
   - 조례에 규정 없는 경우 → 상위법 기준 적용

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 **출력 형식**
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## 🎯 법규 적용 요약

### 📊 핵심 기준 (실질 적용 법규)

| 항목 | 상위법 근거 | 조례 기준 | 실질 적용 |
|------|-------------|-----------|-----------|
| 건폐율 | 건축법 제○조 | OO시 조례 제○조 | ○○% 이하 |
| 용적률 | 국토계획법 제○조 | OO시 조례 제○조 | ○○% 이하 |
| 층수 | ... | ... | ... |

**적용 논리:**
- 건폐율: 국토계획법 제77조에 따라 조례로 위임 → [OO시 도시계획조례 제15조]에 의거, 60% 이하 적용
- 용적률: 국토계획법 제78조에 따라 조례로 위임 → [OO시 도시계획조례 제16조]에 의거, 200% 이하 적용

---

### 📑 전체 관련 법규 목록

#### 1️⃣ 상위법 (국계법)

**건축법**
- 제○조: [조항 내용]
- 제○조: [조항 내용]

**국토의 계획 및 이용에 관한 법률**
- 제○조: [조항 내용]
- 제○조: [조항 내용]

#### 2️⃣ 하위법 (조례)

**OO시 건축 조례**
- 제○조: [조항 내용]
- 제○조: [조항 내용]

**OO시 도시계획 조례**
- 제○조: [조항 내용]
- 제○조: [조항 내용]

---

### 🎯 설계 반영 필수 조항

1. **건폐율: ○○% 이하**
   - 근거: [상위법 조항] + [조례 조항]
   - 설계 시 주의사항: ...

2. **용적률: ○○% 이하**
   - 근거: [상위법 조항] + [조례 조항]
   - 설계 시 주의사항: ...

3. **층수: ○층 이하**
   - 근거: [상위법 조항] + [조례 조항]
   - 설계 시 주의사항: ...

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ **분석 시 주의사항**
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. 반드시 상위법과 조례를 구분하여 표시
2. 조례가 상위법을 강화한 경우 명확히 표시
3. 실질 적용 기준은 조례 우선 원칙 적용
4. 조항 번호를 정확하게 인용
"""
    
    try:
        st.info("⚖️ 법규 위계 분석 중...")
        
        model = genai.GenerativeModel(model_name)
        content_list = gemini_files + [prompt]
        response = model.generate_content(content_list)
        
        st.success("✅ 법규 분석 완료!")
        
        return response.text
        
    except Exception as e:
        st.error(f"❌ 법규 분석 오류: {str(e)}")
        return None


# ================================
# 핵심 함수: 실별 면적표 시각화
# ================================

def visualize_area_table(area_data):
    """
    실별 면적표 시각화
    - Pie Chart: 실별 면적 비중
    - Bar Chart: 전용/공용 면적
    """
    
    if not area_data:
        st.warning("⚠️ 실별 면적표 데이터가 없습니다.")
        return
    
    st.markdown("### 📊 실별 면적표 시각화")
    
    try:
        # 데이터프레임 생성
        df = pd.DataFrame(area_data)
        
        if df.empty or "실명" not in df.columns or "면적" not in df.columns:
            st.warning("⚠️ 면적 데이터 형식이 올바르지 않습니다.")
            return
        
        # 면적을 숫자로 변환
        df["면적"] = pd.to_numeric(df["면적"], errors='coerce')
        df = df.dropna(subset=["면적"])
        
        if df.empty:
            st.warning("⚠️ 유효한 면적 데이터가 없습니다.")
            return
        
        # 1. Pie Chart - 실별 면적 비중
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 🥧 실별 면적 비중")
            
            fig_pie = px.pie(
                df,
                values="면적",
                names="실명",
                title="실별 면적 분포",
                hole=0.3,
                color_discrete_sequence=px.colors.qualitative.Set3
            )
            
            fig_pie.update_traces(
                textposition='inside',
                textinfo='percent+label',
                hovertemplate='<b>%{label}</b><br>면적: %{value}㎡<br>비중: %{percent}<extra></extra>'
            )
            
            fig_pie.update_layout(
                height=400,
                showlegend=True,
                legend=dict(
                    orientation="v",
                    yanchor="middle",
                    y=0.5,
                    xanchor="left",
                    x=1.05
                )
            )
            
            st.plotly_chart(fig_pie, use_container_width=True)
        
        with col2:
            st.markdown("#### 📊 실별 면적 상세")
            
            # 면적 기준 내림차순 정렬
            df_sorted = df.sort_values("면적", ascending=False)
            
            fig_bar = px.bar(
                df_sorted,
                x="실명",
                y="면적",
                title="실별 면적 비교",
                color="면적",
                color_continuous_scale="Blues",
                text="면적"
            )
            
            fig_bar.update_traces(
                texttemplate='%{text}㎡',
                textposition='outside',
                hovertemplate='<b>%{x}</b><br>면적: %{y}㎡<extra></extra>'
            )
            
            fig_bar.update_layout(
                height=400,
                xaxis_title="실명",
                yaxis_title="면적 (㎡)",
                showlegend=False
            )
            
            st.plotly_chart(fig_bar, use_container_width=True)
        
        # 2. 요약 테이블
        st.markdown("#### 📋 면적 요약표")
        
        total_area = df["면적"].sum()
        
        summary_df = df.copy()
        summary_df["비중(%)"] = (summary_df["면적"] / total_area * 100).round(2)
        summary_df = summary_df.sort_values("면적", ascending=False)
        
        # 스타일링
        st.dataframe(
            summary_df,
            use_container_width=True,
            hide_index=True
        )
        
        # 합계
        st.markdown(f"**총 면적: {total_area:,.1f} ㎡**")
        
    except Exception as e:
        st.error(f"❌ 시각화 오류: {str(e)}")


def visualize_private_public_ratio(guideline_data):
    """전용/공용 면적 비율 시각화"""
    
    private_area = guideline_data.get("전용면적", "")
    public_area = guideline_data.get("공용면적", "")
    
    if not private_area or not public_area:
        return
    
    try:
        # 숫자 추출
        private_num = float(re.sub(r'[^0-9.]', '', str(private_area)))
        public_num = float(re.sub(r'[^0-9.]', '', str(public_area)))
        
        if private_num <= 0 or public_num <= 0:
            return
        
        st.markdown("### 🏢 전용/공용 면적 비율")
        
        # Pie Chart
        data = {
            "구분": ["전용면적", "공용면적"],
            "면적": [private_num, public_num]
        }
        
        df = pd.DataFrame(data)
        
        fig = px.pie(
            df,
            values="면적",
            names="구분",
            title="전용/공용 면적 비율",
            color="구분",
            color_discrete_map={"전용면적": "#3b82f6", "공용면적": "#f59e0b"},
            hole=0.4
        )
        
        fig.update_traces(
            textposition='inside',
            textinfo='percent+label+value',
            hovertemplate='<b>%{label}</b><br>면적: %{value}㎡<br>비중: %{percent}<extra></extra>'
        )
        
        fig.update_layout(height=350)
        
        st.plotly_chart(fig, use_container_width=True)
        
        # 비율 계산
        total = private_num + public_num
        private_ratio = (private_num / total * 100)
        public_ratio = (public_num / total * 100)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("전용면적", f"{private_num:,.1f} ㎡", f"{private_ratio:.1f}%")
        
        with col2:
            st.metric("공용면적", f"{public_num:,.1f} ㎡", f"{public_ratio:.1f}%")
        
        with col3:
            st.metric("합계", f"{total:,.1f} ㎡", "100%")
        
    except Exception as e:
        st.error(f"❌ 전용/공용 비율 시각화 오류: {str(e)}")


# ================================
# 핵심 함수: 보고서 생성
# ================================

def set_cell_background(cell, color):
    """표 셀 배경색 설정"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_professional_report(guideline_data, regulation_result, area_data):
    """
    전문 보고서 생성 (상위법/하위법 구분)
    """
    
    doc = Document()
    
    # 페이지 여백
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 제목
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("건축 공모 법규 검토서")
    title_run.font.name = '맑은 고딕'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{guideline_data.get('프로젝트명', '프로젝트')}")
    subtitle_run.font.name = '맑은 고딕'
    subtitle_run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 1. 공모 개요
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    section1 = doc.add_paragraph()
    section1_run = section1.add_run("1. 공모 개요")
    section1_run.font.name = '맑은 고딕'
    section1_run.font.size = Pt(14)
    section1_run.font.bold = True
    
    # 개요 표
    summary_table = doc.add_table(rows=8, cols=2)
    summary_table.style = 'Table Grid'
    
    # 헤더
    headers = ['항목', '내용']
    header_cells = summary_table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        set_cell_background(cell, 'D3D3D3')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(11)
        run.font.bold = True
    
    # 데이터
    current_date = datetime.now().strftime("%Y년 %m월 %d일")
    
    data_rows = [
        ('프로젝트명', guideline_data.get('프로젝트명', '-')),
        ('위치', guideline_data.get('위치', '-')),
        ('지역지구', guideline_data.get('지역지구', '-')),
        ('용도', guideline_data.get('용도', '-')),
        ('대지면적', guideline_data.get('대지면적', '-')),
        ('건폐율', guideline_data.get('건폐율', '-')),
        ('용적률', guideline_data.get('용적률', '-'))
    ]
    
    for idx, (label, value) in enumerate(data_rows, 1):
        row_cells = summary_table.rows[idx].cells
        
        p0 = row_cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run0 = p0.add_run(label)
        run0.font.name = '맑은 고딕'
        run0.font.size = Pt(10)
        run0.font.bold = True
        
        p1 = row_cells[1].paragraphs[0]
        run1 = p1.add_run(str(value))
        run1.font.name = '맑은 고딕'
        run1.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 2. 법규 적용 기준 (상위법/하위법)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    section2 = doc.add_paragraph()
    section2_run = section2.add_run("2. 법규 적용 기준")
    section2_run.font.name = '맑은 고딕'
    section2_run.font.size = Pt(14)
    section2_run.font.bold = True
    
    # 법규 표 (4열: 항목, 상위법 근거, 조례 기준, 실질 적용)
    reg_table = doc.add_table(rows=4, cols=4)
    reg_table.style = 'Table Grid'
    
    # 헤더
    reg_headers = ['항목', '법적 근거 (상위법)', '실무 적용 기준 (조례)', '최종 적용']
    header_cells = reg_table.rows[0].cells
    for idx, header in enumerate(reg_headers):
        cell = header_cells[idx]
        set_cell_background(cell, 'D3D3D3')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)
        run.font.bold = True
    
    # 데이터 (예시 - 실제로는 regulation_result에서 파싱)
    reg_data = [
        ('건폐율', '국토계획법 제77조', 'OO시 조례 제15조', '60% 이하'),
        ('용적률', '국토계획법 제78조', 'OO시 조례 제16조', '200% 이하'),
        ('층수', '건축법 제60조', 'OO시 조례 제20조', '10층 이하')
    ]
    
    for idx, (item, upper, lower, final) in enumerate(reg_data, 1):
        row_cells = reg_table.rows[idx].cells
        
        for col_idx, text in enumerate([item, upper, lower, final]):
            p = row_cells[col_idx].paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 3. 상세 분석 결과
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    section3 = doc.add_paragraph()
    section3_run = section3.add_run("3. 상세 분석 결과")
    section3_run.font.name = '맑은 고딕'
    section3_run.font.size = Pt(14)
    section3_run.font.bold = True
    
    detail_para = doc.add_paragraph()
    detail_run = detail_para.add_run(regulation_result[:2000] if regulation_result else "분석 결과 없음")
    detail_run.font.name = '맑은 고딕'
    detail_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 4. 실별 면적표
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    if area_data:
        section4 = doc.add_paragraph()
        section4_run = section4.add_run("4. 실별 면적표")
        section4_run.font.name = '맑은 고딕'
        section4_run.font.size = Pt(14)
        section4_run.font.bold = True
        
        area_table = doc.add_table(rows=len(area_data)+1, cols=2)
        area_table.style = 'Table Grid'
        
        # 헤더
        area_headers = ['실명', '면적 (㎡)']
        header_cells = area_table.rows[0].cells
        for idx, header in enumerate(area_headers):
            cell = header_cells[idx]
            set_cell_background(cell, 'D3D3D3')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10)
            run.font.bold = True
        
        # 데이터
        for idx, item in enumerate(area_data, 1):
            row_cells = area_table.rows[idx].cells
            
            p0 = row_cells[0].paragraphs[0]
            run0 = p0.add_run(item.get('실명', '-'))
            run0.font.name = '맑은 고딕'
            run0.font.size = Pt(9)
            
            p1 = row_cells[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run1 = p1.add_run(str(item.get('면적', '-')))
            run1.font.name = '맑은 고딕'
            run1.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 법적 고지
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    doc.add_page_break()
    
    notice_title = doc.add_paragraph()
    notice_title_run = notice_title.add_run("⚖️ 법적 고지")
    notice_title_run.font.name = '맑은 고딕'
    notice_title_run.font.size = Pt(12)
    notice_title_run.font.bold = True
    
    notice_text = doc.add_paragraph()
    notice_run = notice_text.add_run(
        "본 법규 검토서는 AI 기반 분석 도구를 활용하여 작성된 참고 자료입니다.\n\n"
        f"작성일: {current_date}\n"
        "작성: 건축 공모 & 법규 분석 시스템 v4.0\n\n"
        "All intellectual property rights belong to Kim Doyoung."
    )
    notice_run.font.name = '맑은 고딕'
    notice_run.font.size = Pt(9)
    notice_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # 메모리에 저장
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    return docx_file


# ================================
# 메인 UI
# ================================

st.markdown(
    '<div class="main-title">🏛️ 건축 공모 & 법규 분석 시스템'
    '<span class="version-badge">v4.0</span></div>',
    unsafe_allow_html=True
)

st.markdown("""
<div style='text-align: center; margin-bottom: 2rem;'>
    <p style='font-size: 1.1rem; color: #555;'>
        🚀 <b>Gemini 2.5 Flash</b> 기반 전문 분석 시스템<br>
        공모지침서 + 법규 교차분석 | 법규 위계 분석 | 실별 면적표 시각화
    </p>
</div>
""", unsafe_allow_html=True)

st.divider()

# ================================
# A. 공모지침서 업로드
# ================================
st.markdown('<div class="section-header"><h2>📄 A. 공모지침서 업로드 (단일)</h2></div>', 
           unsafe_allow_html=True)

competition_file = st.file_uploader(
    "공모지침서 PDF 파일을 선택하세요",
    type=['pdf'],
    help="건축 공모 지침서 1개 파일"
)

if competition_file:
    st.success(f"✅ {competition_file.name} ({competition_file.size / 1024:.1f} KB)")

st.divider()

# ================================
# B. 법규 PDF 업로드
# ================================
st.markdown('<div class="section-header"><h2>⚖️ B. 관련 법규 PDF 업로드 (다중)</h2></div>', 
           unsafe_allow_html=True)

regulation_files = st.file_uploader(
    "법규 PDF 파일들을 선택하세요 (여러 개 가능)",
    type=['pdf'],
    accept_multiple_files=True,
    help="건축법, 조례 등 관련 법규 문서"
)

if regulation_files:
    st.success(f"✅ {len(regulation_files)}개 파일 선택됨")
    
    for idx, file in enumerate(regulation_files, 1):
        st.markdown(f"**{idx}.** 📄 {file.name} (`{file.size / 1024:.1f} KB`)")

st.divider()

# ================================
# 분석 실행
# ================================
st.markdown('<div class="section-header"><h2>🔍 분석 실행</h2></div>', 
           unsafe_allow_html=True)

analyze_button = st.button(
    "🚀 통합 분석 시작",
    type="primary",
    use_container_width=True
)

if analyze_button:
    if not api_key:
        st.error("❌ API 키를 먼저 설정하세요!")
    elif not competition_file:
        st.error("❌ 공모지침서를 업로드하세요!")
    elif not regulation_files:
        st.error("❌ 법규 PDF를 최소 1개 이상 업로드하세요!")
    else:
        st.markdown("---")
        st.markdown("### 🚀 분석 진행 중...")
        
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # 1단계: 공모지침서 업로드 및 분석
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        st.markdown("#### 📄 1단계: 공모지침서 분석")
        
        comp_gemini_file = upload_pdf_to_gemini(competition_file, "공모지침서")
        
        if comp_gemini_file:
            guideline_data = analyze_competition_guidelines(comp_gemini_file, selected_model)
            
            if guideline_data:
                # 세션 상태 저장
                st.session_state['guideline_data'] = guideline_data
                
                # 공모 개요 표시
                st.markdown("##### 🎯 공모 개요 요약")
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("프로젝트", guideline_data.get('프로젝트명', '-'))
                    st.metric("위치", guideline_data.get('위치', '-'))
                
                with col2:
                    st.metric("지역지구", guideline_data.get('지역지구', '-'))
                    st.metric("용도", guideline_data.get('용도', '-'))
                
                with col3:
                    st.metric("건폐율", guideline_data.get('건폐율', '-'))
                    st.metric("용적률", guideline_data.get('용적률', '-'))
                
                st.divider()
                
                # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                # 2단계: 법규 업로드 및 분석
                # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                st.markdown("#### ⚖️ 2단계: 법규 분석 (위계 포함)")
                
                reg_gemini_files = upload_multiple_pdfs(regulation_files, "법규")
                
                if reg_gemini_files:
                    regulation_result = analyze_regulations_with_hierarchy(
                        reg_gemini_files,
                        guideline_data,
                        selected_model
                    )
                    
                    if regulation_result:
                        st.session_state['regulation_result'] = regulation_result
                        
                        # 법규 분석 결과 표시
                        st.markdown("##### 📋 법규 분석 결과")
                        st.markdown(regulation_result)
                        
                        st.divider()
                        
                        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                        # 3단계: 실별 면적표 시각화
                        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                        if include_visualization:
                            st.markdown("#### 📊 3단계: 실별 면적표 시각화")
                            
                            area_data = guideline_data.get('실별면적표', [])
                            
                            if area_data:
                                visualize_area_table(area_data)
                                visualize_private_public_ratio(guideline_data)
                            else:
                                st.info("ℹ️ 실별 면적표 데이터가 지침서에 없습니다.")
                        
                        st.divider()
                        
                        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                        # 보고서 다운로드
                        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                        st.markdown("### 💾 결과 저장")
                        
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            # Markdown 다운로드
                            combined_result = f"""# 공모 분석 결과

## 공모 개요
{json.dumps(guideline_data, ensure_ascii=False, indent=2)}

## 법규 분석
{regulation_result}
"""
                            
                            st.download_button(
                                label="📝 Markdown 다운로드",
                                data=combined_result,
                                file_name=f"공모분석_{guideline_data.get('프로젝트명', 'project')}_{datetime.now().strftime('%Y%m%d')}.md",
                                mime="text/markdown",
                                use_container_width=True
                            )
                        
                        with col2:
                            # 전문 보고서 다운로드
                            try:
                                report = create_professional_report(
                                    guideline_data,
                                    regulation_result,
                                    guideline_data.get('실별면적표', [])
                                )
                                
                                st.download_button(
                                    label="📄 전문 보고서 (docx)",
                                    data=report,
                                    file_name=f"법규검토서_{guideline_data.get('프로젝트명', 'project')}_{datetime.now().strftime('%Y%m%d')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True,
                                    type="primary"
                                )
                            except Exception as e:
                                st.error(f"보고서 생성 오류: {str(e)}")
                        
                        with col3:
                            # JSON 다운로드
                            json_data = {
                                "공모개요": guideline_data,
                                "법규분석": regulation_result
                            }
                            
                            st.download_button(
                                label="📊 JSON 다운로드",
                                data=json.dumps(json_data, ensure_ascii=False, indent=2),
                                file_name=f"분석데이터_{datetime.now().strftime('%Y%m%d')}.json",
                                mime="application/json",
                                use_container_width=True
                            )

# 푸터
st.divider()
st.markdown(f"""
<div style='text-align: center; color: gray; padding: 20px;'>
    <small>
    🚀 <b>Powered by Google Gemini 2.5 Flash</b><br>
    건축 공모 & 법규 분석 시스템 v4.0 Professional Edition<br><br>
    ⚖️ <b>법적 고지:</b> 본 분석은 AI 기반 참고 자료이며, 법적 효력이 없습니다.<br>
    실제 설계 시 반드시 전문가의 검토를 받으시기 바랍니다.<br><br>
    <b>Version 4.0</b> | 2026년 2월<br>
    All intellectual property rights belong to Kim Doyoung.
    </small>
</div>
""", unsafe_allow_html=True)